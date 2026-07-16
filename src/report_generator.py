"""
LLM 报告生成模块 V2.0

峤远主脑协议驱动的双版本报告生成器：
- 专业版（7模块，面向财务总监/财务人员）
- 大白话版（3段式，面向老板/非财务背景决策者）

核心特性：
1. 主脑协议六条戒律驱动分析框架
2. 三模型数据注入（杜邦/Z-Score/F-Score）
3. 增强指标按Q2选择动态注入
4. 模型解释固化文本（不走LLM）
5. Council多角色机制（触发条件：融资能力/Z<1.81）
6. 行业基准对标
7. 优化改进建议四段式
"""

import os
import io
import re
import json
import logging
from typing import Optional
from math import isnan

_logger = logging.getLogger(__name__)

# ═══════════════════════════════════════════════════════════
# 主脑协议 System Prompts
# ═══════════════════════════════════════════════════════════

_PROFESSIONAL_SYSTEM_PROMPT = """你是峤远，企业财务总监级AI专家。你的读者是财务总监、财务经理或银行信贷员，他们有专业基础，需要通过报告验证自己的判断。

你拥有四栈专家能力：财务分析+税务筹划+融资顾问+法务合规。

你必须遵守六条戒律：
1. 数据真实优先：所有分析必须基于提供的指标数据，不编造数据
2. 风险穿透：不停留在表面指标，穿透到业务实质
3. 行业基准锚定：指标评价必须对标行业基准
4. 趋势研判从严：单期数据不做趋势判断
5. 客户视角校准：分析维度匹配客户关注点
6. 可执行性统领：建议必须具体可操作，禁止空话

输出要求：
1. 使用标准财务术语，保留专业表达
2. 每个指标展示：数值+行业基准+偏差幅度+评级
3. 模型分析展示完整分解树/分变量/信号明细
4. 优化建议按"问题→原因→措施→预期效果"四段式
5. 引用六条戒律作为分析逻辑依据
6. 语气客观、专业、数据驱动
7. 直接输出报告正文，不要以"好的"开头"""

_PLAIN_SYSTEM_PROMPT = """你是峤远，一个专门给老板讲财务的AI助手。你的读者是没学过财务的小老板和个体户。

你必须遵守八条铁律：
1. 每个专业术语后面必须跟一句大白话翻译
   正确："流动比率0.85（手里能很快变现的钱只够还85%的短期欠款）"
   错误："流动比率0.85，低于行业基准"
2. 不用公式，只说结论和含义
3. 每个问题用"四问"展开：有没有问题？严不严重？什么原因？不管会怎样？
4. 每个建议用"四答"展开：做什么？为什么？怎么做？多久见效？
5. 用红黄绿灯代替数字评级
6. 禁止出现：同比、环比、权益乘数、EBITDA、边际贡献等术语（如必须使用，立即在括号内翻译）
7. 结论先行：每段第一句话就是结论
8. 语气像一个懂行的老朋友在聊天，不像在念报告
9. 称呼读者统一用"老板"或"您好"，禁止编造人名（如"老张""老李"等）

直接输出报告正文，不要以"好的"开头。"""

# ═══════════════════════════════════════════════════════════
# 模型解释固化文本（不走LLM，直接拼接）
# ═══════════════════════════════════════════════════════════

_MODEL_EXPLANATION_DUPONT = """━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
📌 模型一：杜邦分析 — 你的赚钱效率拆开来看
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

【这是什么？】
杜邦分析是杜邦公司发明的一种财务分解方法，把"净资产收益率（ROE）"——也就是你每投入1块钱本金能赚多少——拆成三个零件，让你看清赚钱的来源。

ROE = 净利率 × 资产周转率 × 权益乘数

把这三个零件翻译成大白话：
  • 净利率 = 每卖100块钱东西，到手能赚多少利润 → 衡量"卖得贵不贵"
  • 资产周转率 = 你的资产一年能转几圈 → 衡量"卖得快不快"
  • 权益乘数 = 你借了多少钱来做生意 → 衡量"杠杆大不大"

【怎么读？】
看分解树图，三个零件哪个高哪个低，就知道你的ROE主要靠什么驱动：
  • 如果靠净利率高 → 你走的是"高利润率路线"（产品有定价权）
  • 如果靠周转率高 → 你走的是"薄利多销路线"（靠规模和效率）
  • 如果靠权益乘数高 → 你走的是"杠杆路线"（靠借钱放大收益）
    ⚠️ 注意：杠杆是双刃剑，赚钱时放大收益，亏钱时也放大亏损

【你应该关注什么？】
1. 三个零件哪个最强、哪个最弱 → 知道该补哪块
2. 最弱的那个就是你的"提升杠杆点" → 优化它对ROE的提升效果最大
3. 如果权益乘数过高（比如>3）→ 赚钱靠借钱，风险在积聚
4. 如果有上年数据 → 看哪个零件变了，是变好的还是变坏的"""

_MODEL_EXPLANATION_ZSCORE = """━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
📌 模型二：Altman Z-Score — 你的企业离破产有多远？
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

【这是什么？】
Z-Score是纽约大学Edward Altman教授发明的破产预测模型，用5个财务指标算出一个综合分数，被全球银行和投资机构广泛使用。它的准确率在预测2年内企业是否破产方面超过80%。

【怎么读？】
看仪表盘上你的Z值落在哪个区域：
  🟢 Z > 2.99  → 安全区：企业财务健康，短期无破产风险
  🟡 1.81-2.99 → 灰色区：有一定风险，需要关注薄弱环节
  🔴 Z < 1.81  → 危险区：破产风险较高，需要立即采取行动

【你应该关注什么？】
1. Z值落在哪个区域 → 整体风险定位
2. 5个分变量哪个拖了后腿 → 那个就是最需要改善的
   X1 营运资金/总资产 → 短期周转能力
   X2 留存收益/总资产 → 历史积累能力
   X3 EBIT/总资产 → 经营盈利能力
   X4 权益/总负债 → 资本结构稳健性
   X5 营业收入/总资产 → 资产产出效率
3. 如果在灰色区或危险区 → 必须优先改善最弱的1-2个分变量
4. Z值不是越高越好 → 过高的Z值可能意味着资产闲置、未充分利用杠杆"""

_MODEL_EXPLANATION_FSCORE = """━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
📌 模型三：Piotroski F-Score — 你的企业在走上坡路还是下坡路？
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

【这是什么？】
F-Score是芝加哥大学Joseph Piotroski教授发明的财务质量评分模型，用9个"是/否"信号来衡量企业财务状况是在变好还是变差。每满足一项得1分，满分9分。它不看你"现在好不好"，而是看你"方向对不对"。

【怎么读？】
看你的总分落在哪个区间：
  🟢 7-9分 → 财务状况持续改善，走上坡路
  🟡 4-6分 → 财务状况基本平稳，没有明显改善或恶化
  🔴 0-3分 → 财务状况在恶化，走下坡路，需要警惕

9项信号分三大块：
  【盈利能力】F1 有没有赚钱（ROA>0）
              F2 有没有真金白银（经营现金流>0）
              F3 比去年赚得多吗（ROA同比提升）
              F4 赚的钱是不是真钱（现金流>净利润）
  【杠杆与流动性】F5 负债有没有减少
                  F6 短期偿债能力有没有变强
                  F7 有没有稀释股东权益（增发新股）
  【营运效率】F8 毛利率有没有提升
              F9 资产周转有没有变快

【你应该关注什么？】
1. 总分落在哪个区间 → 你的企业"方向感"如何
2. 失分的项在哪 → 那些就是正在变差的方面
3. F3/F5/F6/F7/F8/F9需要上年数据才能算 → 单期只能评0-3分
   ⚠️ 如果只上传了1期报表，F-Score只有部分评分（满分3分），建议上传上年报表获取完整9分评分
4. F4（现金流>净利润）失分要特别警惕 → 账上有利润但收不回钱，是中小企业最常见的"纸面富贵"陷阱
5. F-Score和Z-Score要一起看：
   Z-Score看"你现在安不安全"，F-Score看"你在往哪走"
   Z值高+F-Score高 = 最理想：安全且在变好
   Z值高+F-Score低 = 需警惕：现在还行但在走下坡路
   Z值低+F-Score高 = 有希望：虽然不安全但在改善
   Z值低+F-Score低 = 最危险：不安全且还在恶化"""

# ═══════════════════════════════════════════════════════════
# 免责声明
# ═══════════════════════════════════════════════════════════

_DISCLAIMER_PRO = """**免责与合规声明**

1. 本报告由霖信莯咨询·F-Analyzer系统（峤远主脑协议V1.0驱动）自动生成，不构成专业财务意见。
2. 所有指标数值由财务计算引擎确定，LLM仅负责解读，可能存在偏差。
3. 行业基准数据来源于国家统计局中小企业均值+A股上市公司公开数据，与具体企业存在差异，仅供参考。
4. 杜邦分析、Altman Z-Score、Piotroski F-Score三个模型均为学术模型，参数基于公开研究，不保证对特定企业的预测准确性。
5. 重大财务决策请咨询持牌专业顾问，霖信莯咨询可提供人工复核服务。

---
**财务分析辅助工具，输出不构成正式财务意见或融资承诺**"""

_DISCLAIMER_PLAIN = """---
📋 声明：这份报告是AI自动生成的，里面的数字是算出来的（准），但解读是AI写的（可能不完全准）。如果你要做重大决定（比如贷款、投资），建议找个有证的会计师再确认一下。"""


class ReportGenerator:
    """
    峤远主脑协议 V1.0 驱动的双版本报告生成器
    """

    def __init__(
        self,
        api_key: Optional[str] = None,
        base_url: str = "https://api.deepseek.com",
        model: str = "deepseek-chat",
    ):
        self._api_key = api_key or os.getenv("DEEPSEEK_API_KEY", "")
        self._base_url = base_url
        self._model = model
        self._client = None

    def _ensure_client(self):
        if self._client is None and self._api_key:
            from openai import OpenAI
            self._client = OpenAI(api_key=self._api_key, base_url=self._base_url)

    # ═══════════════════════════════════════════════════════════
    # 主入口
    # ═══════════════════════════════════════════════════════════

    def generate(
        self,
        ratios: dict,
        mapped_df=None,
        check_results=None,
        analysis_goal: str = "财务健康检查",
        report_period: str = "未指定报告期",
        calibration: dict = None,
        enhanced_metrics: dict = None,
        model_results: dict = None,
        prior_ratios: dict = None,
        growth_metrics: dict = None,
        industry: str = "其他",
    ) -> dict:
        """
        生成双版本报告

        Args:
            ratios: 基础指标计算结果
            mapped_df: 标准化科目DataFrame
            check_results: 质量校验结果
            analysis_goal: 分析目标
            report_period: 报告期
            calibration: 选择题校准结果 {q1_industry, q2_dimensions, q3_purpose}
            enhanced_metrics: 增强指标计算结果
            model_results: 三模型分析结果 {dupont, zscore, fscore}
            prior_ratios: 上期指标（如果有）
            growth_metrics: 成长能力指标（如果有上期数据）
            industry: 行业

        Returns:
            {
                "professional": {模块名: 内容},
                "plain": str,
                "model_explanations": str,  # 固化文本
            }
        """
        calibration = calibration or {}
        model_results = model_results or {}

        # 构建共享数据上下文
        context = self._build_shared_context(
            ratios, mapped_df, check_results, analysis_goal, report_period,
            calibration, enhanced_metrics, model_results, prior_ratios, growth_metrics, industry
        )

        # 模型解释固化文本（两个版本共用，但展示方式不同）
        model_explanations = self._build_model_explanations()

        # 生成专业版
        professional = self._generate_professional(context)

        # 生成大白话版
        plain = self._generate_plain(context)

        return {
            "professional": professional,
            "plain": plain,
            "model_explanations": model_explanations,
            "analysis_goal": analysis_goal,
            "report_period": report_period,
        }

    # ═══════════════════════════════════════════════════════════
    # 共享数据上下文构建
    # ═══════════════════════════════════════════════════════════

    def _build_shared_context(
        self, ratios, mapped_df, check_results, analysis_goal, report_period,
        calibration, enhanced_metrics, model_results, prior_ratios, growth_metrics, industry
    ) -> dict:
        """构建两个版本共享的数据上下文"""
        from src.industry_benchmarks import get_industry_benchmarks
        benchmarks = get_industry_benchmarks(industry)

        # 关键科目
        key_accounts = {}
        if mapped_df is not None:
            for col in ["total_assets", "total_liabilities", "total_equity",
                        "revenue", "net_profit", "cash_and_equivalents",
                        "operating_cash_flow", "inventory", "accounts_receivable"]:
                if col in mapped_df.columns:
                    val = mapped_df[col].iloc[0]
                    if not (isinstance(val, float) and isnan(val)):
                        key_accounts[col] = float(val)

        # 指标格式化文本
        ratios_text = self._all_ratios_to_text(ratios, benchmarks)
        enhanced_text = self._enhanced_to_text(enhanced_metrics, benchmarks) if enhanced_metrics else ""
        quality_notes = self._make_quality_notes(check_results) if check_results else "暂无质量校验数据"

        # 模型数据
        dupont_text = model_results.get("dupont_summary", "")
        zscore_text = model_results.get("zscore_summary", "")
        fscore_text = model_results.get("fscore_summary", "")

        # 成长指标
        growth_text = self._growth_to_text(growth_metrics) if growth_metrics else ""

        # Council 触发判断
        z_val = ratios.get("财务预警", {}).get("Altman Z-score", {}).get("value", 0)
        trigger_council = False
        if "融资能力" in calibration.get("q2_dimensions", []):
            trigger_council = True
        if isinstance(z_val, float) and z_val < 1.81:
            trigger_council = True

        return {
            "ratios": ratios,
            "ratios_text": ratios_text,
            "enhanced_text": enhanced_text,
            "key_accounts": key_accounts,
            "quality_notes": quality_notes,
            "benchmarks": benchmarks,
            "industry": industry,
            "calibration": calibration,
            "analysis_goal": analysis_goal,
            "report_period": report_period,
            "dupont_text": dupont_text,
            "zscore_text": zscore_text,
            "fscore_text": fscore_text,
            "growth_text": growth_text,
            "trigger_council": trigger_council,
            "has_prior": prior_ratios is not None,
        }

    # ═══════════════════════════════════════════════════════════
    # 专业版报告生成
    # ═══════════════════════════════════════════════════════════

    def _generate_professional(self, ctx: dict) -> dict:
        """生成专业版7模块报告"""
        report = {
            "执行摘要": None,
            "财务诊断": None,
            "经典模型分析": None,
            "优化改进建议": None,
            "税务分析": None,
            "融资能力评估": None,
            "免责声明": _DISCLAIMER_PRO,
        }

        if not self._api_key:
            for k in report:
                if k != "免责声明":
                    report[k] = f"[{k}]\n本模块需API连接生成。当前展示指标计算值。\n"
            # 至少把模型解释拼进去
            report["经典模型分析"] = self._build_model_explanations() + "\n\n" + ctx["dupont_text"] + "\n" + ctx["zscore_text"] + "\n" + ctx["fscore_text"]
            return report

        self._ensure_client()

        # 模块1：执行摘要
        report["执行摘要"] = self._call_llm(
            _PROFESSIONAL_SYSTEM_PROMPT,
            self._build_summary_prompt(ctx)
        )

        # 模块2：财务诊断
        report["财务诊断"] = self._call_llm(
            _PROFESSIONAL_SYSTEM_PROMPT,
            self._build_diagnosis_prompt(ctx)
        )

        # 模块3：经典模型分析（固化文本 + LLM联立解读）
        model_section = self._build_model_explanations()
        model_section += "\n\n---\n\n**您的三模型分析数据：**\n\n"
        model_section += ctx["dupont_text"] + "\n\n"
        model_section += ctx["zscore_text"] + "\n\n"
        model_section += ctx["fscore_text"] + "\n\n"
        model_section += self._call_llm(
            _PROFESSIONAL_SYSTEM_PROMPT,
            self._build_model_synthesis_prompt(ctx)
        )
        report["经典模型分析"] = model_section

        # 模块4：优化改进建议
        report["优化改进建议"] = self._call_llm(
            _PROFESSIONAL_SYSTEM_PROMPT,
            self._build_optimization_prompt(ctx)
        )

        # 模块5：税务分析
        report["税务分析"] = self._call_llm(
            _PROFESSIONAL_SYSTEM_PROMPT,
            self._build_tax_prompt(ctx)
        )

        # 模块6：融资能力评估（Council触发时）
        if ctx["trigger_council"]:
            report["融资能力评估"] = self._call_llm(
                _PROFESSIONAL_SYSTEM_PROMPT,
                self._build_council_prompt(ctx)
            )

        return report

    def _build_summary_prompt(self, ctx: dict) -> str:
        cal = ctx["calibration"]
        q2 = ", ".join(cal.get("q2_dimensions", ["标准版（未校准）"]))
        q3 = cal.get("q3_purpose", "总体体检")
        ka = ctx["key_accounts"]
        ka_text = "\n".join(f"  {k}: {v:,.0f}" for k, v in ka.items()) if ka else "暂无"

        return f"""请生成执行摘要（≥400字）。

**企业信息**
报告期: {ctx['report_period']}
行业: {ctx['industry']}
分析目标: {ctx['analysis_goal']}
客户关注维度: {q2}
报告用途: {q3}

**关键科目**
{ka_text}

**核心指标速览**
{ctx['ratios_text'][:800]}

**数据质量**
{ctx['quality_notes']}

**撰写要求**
1. 第一段（≥150字）：企业财务状况总体判断+总体评级（优秀/良好/一般/偏弱/危险）
2. 第二段（≥150字）：关键指标一览表（用markdown表格，含指标名/数值/行业基准/偏差/评级）
3. 第三段（≥100字）：核心风险提示+核心优势亮点
4. 如果有数据质量异常，首句提示
5. 遵守戒律1（数据真实优先）和戒律3（行业基准锚定）
"""

    def _build_diagnosis_prompt(self, ctx: dict) -> str:
        cal = ctx["calibration"]
        q2 = cal.get("q2_dimensions", [])
        q2_text = "、".join(q2) if q2 else "标准版（全维度均衡）"

        enhanced_section = ""
        if ctx["enhanced_text"]:
            enhanced_section = f"""
**增强指标（根据客户关注维度抽取）**
客户选择了：{q2_text}

{ctx['enhanced_text']}
"""

        growth_section = ""
        if ctx["growth_text"]:
            growth_section = f"""
**成长能力指标（双期数据）**
{ctx['growth_text']}
"""
        else:
            growth_section = "\n注：仅上传了1期报表，成长能力指标无法计算（戒律4：趋势研判从严）。\n"

        return f"""请生成财务诊断模块（≥1500字）。

**行业基准（{ctx['industry']}）**
{self._benchmarks_to_text(ctx['benchmarks'])}

**基础指标**
{ctx['ratios_text']}
{enhanced_section}{growth_section}
**数据质量**
{ctx['quality_notes']}

**客户关注维度**: {q2_text}

**撰写要求**
1. 按七大类（偿债/盈利/营运/成长/现金流/税务/预警）逐项诊断
2. 每个指标展示：数值+行业基准+偏差幅度+评级（优/良/中/差）
3. 对客户选择的关注维度（{q2_text}）做重点深度分析，其他维度简述
4. 戒律2（风险穿透）：每个异常指标必须追溯到业务原因
5. 戒律4（趋势研判从严）：单期数据不做趋势判断
6. 戒律3（行业基准锚定）：所有评级对标{ctx['industry']}行业基准
"""

    def _build_model_synthesis_prompt(self, ctx: dict) -> str:
        return f"""请基于客户实际数据，对三个模型进行联立解读（≥800字）。

**杜邦分析数据**
{ctx['dupont_text']}

**Z-Score数据**
{ctx['zscore_text']}

**F-Score数据**
{ctx['fscore_text']}

**行业**: {ctx['industry']}

**撰写要求**
1. 【杜邦分析】客户ROE数值→三因素分解→主要驱动因子→改进建议
   格式：你的ROE是X%，拆开来看：净利率X%、资产周转率X次、权益乘数X倍。主要靠XX驱动，XX是短板。
2. 【Z-Score】客户Z值→所处区域→5个分变量中最弱的→改善方向
3. 【F-Score】客户F值→所处区间→失分项→警示（特别注意F4盈余质量）
4. 【三模型联立】综合判断+Z×F矩阵定位+行动建议
5. 如有双期数据，分析趋势变化方向
"""

    def _build_optimization_prompt(self, ctx: dict) -> str:
        cal = ctx["calibration"]
        q2 = "、".join(cal.get("q2_dimensions", []))
        q3 = cal.get("q3_purpose", "总体体检")

        return f"""请生成优化改进建议模块（≥1000字）。

**全部指标数据**
{ctx['ratios_text']}

{ctx['enhanced_text'] if ctx['enhanced_text'] else ''}

**数据质量**
{ctx['quality_notes']}

**客户关注维度**: {q2}
**报告用途**: {q3}

**撰写要求**
按四大维度给出建议，每条遵循"问题→原因→措施→预期效果"四段式：

1. 【报表优化】如有数据质量问题（负值/极值/不平衡），给出报表修正建议
2. 【财务管理】针对偏离基准的指标，给出改善建议
3. 【成本管控】如有成本费用率数据，给出降本增效建议
4. 【税收筹划】分析税负率，给出税务优化建议

戒律6（可执行性统领）：
- ✅ "建议将客户信用账期从120天缩短至60天" 
- ❌ "建议加强应收账款管理"
每条建议必须具体到操作步骤，并给出预期效果（定量或半定量）。

如果客户选择了特定关注维度（{q2}），该维度建议需更详细（≥300字）。
"""

    def _build_tax_prompt(self, ctx: dict) -> str:
        ratios = ctx["ratios"]
        etr = ratios.get("税务指标", {}).get("实际税率", {}).get("value", "N/A")
        ctb = ratios.get("税务指标", {}).get("综合税负率", {}).get("value", "N/A")
        gm = ratios.get("盈利能力", {}).get("毛利率", {}).get("value", "N/A")
        bench_etr = ctx["benchmarks"].get("实际税率", "N/A")
        bench_ctb = ctx["benchmarks"].get("综合税负率", "N/A")

        return f"""请生成税务分析模块（≥400字）。

**税务核心指标**
实际税率: {etr*100:.2f}% if isinstance(etr, float) else {etr}
综合税负率: {ctb*100:.2f}% if isinstance(ctb, float) else {ctb}
毛利率: {gm*100:.2f}% if isinstance(gm, float) else {gm}

**行业基准（{ctx['industry']}）**
实际税率基准: {bench_etr*100:.2f}% if isinstance(bench_etr, float) else {bench_etr}
综合税负率基准: {bench_ctb*100:.2f}% if isinstance(bench_ctb, float) else {bench_ctb}

**撰写要求**
1. 实际税率偏离25%标准时的原因分析（小微企业优惠/亏损弥补/虚列费用等）
2. 综合税负率过低→被税务机关质疑的风险
3. 综合税负率过高→税收优化空间
4. 2-4条明确的风险提示和合规建议
5. 戒律2（风险穿透）：穿透到业务实质
"""

    def _build_council_prompt(self, ctx: dict) -> str:
        z_val = ctx["ratios"].get("财务预警", {}).get("Altman Z-score", {}).get("value", 0)
        return f"""请以Council多角色模式生成融资能力评估模块（≥500字）。

**触发原因**: {"客户选择融资能力评估" if "融资能力" in ctx['calibration'].get('q2_dimensions', []) else f"Altman Z-Score={z_val:.2f}<1.81"}

**核心指标**
{ctx['ratios_text'][:600]}

**Council六角色规则**：
每个角色≤200字，串行输出，必须暴露分歧：
1. 【CFO视角】整体财务健康判断+融资可行性
2. 【税务师视角】税务合规对融资的影响
3. 【融资顾问视角】银行信贷可贷性评估+建议融资方案
4. 【内控审计视角】内控风险对融资的阻碍
5. 【成本管控视角】成本结构对还款能力的影响
6. 【反对者视角】挑战上述乐观判断，暴露隐患

最后给出：分歧点+综合结论+行动建议。
"""

    # ═══════════════════════════════════════════════════════════
    # 大白话版报告生成
    # ═══════════════════════════════════════════════════════════

    def _generate_plain(self, ctx: dict) -> str:
        """生成大白话版3段式报告"""
        if not self._api_key:
            return "本报告需API连接生成。\n\n" + _DISCLAIMER_PLAIN

        self._ensure_client()
        return self._call_llm(
            _PLAIN_SYSTEM_PROMPT,
            self._build_plain_prompt(ctx)
        )

    def _build_plain_prompt(self, ctx: dict) -> str:
        cal = ctx["calibration"]
        q2 = "、".join(cal.get("q2_dimensions", []))
        q3 = cal.get("q3_purpose", "总体体检")
        ka = ctx["key_accounts"]
        ka_text = "\n".join(f"  {k}: {v:,.0f}" for k, v in ka.items()) if ka else "暂无"

        # 简化模型数据
        dupont = ctx.get("dupont_text", "")
        zscore = ctx.get("zscore_text", "")
        fscore = ctx.get("fscore_text", "")

        return f"""请生成大白话版财务分析报告（共3段，总字数2000-3500字）。

**企业信息**
报告期: {ctx['report_period']}
行业: {ctx['industry']}
客户关注: {q2}
报告用途: {q3}

**关键数据（这些数字是算出来的，你可以放心引用）**
{ka_text}

**全部指标（请只引用数字，不要用术语名）**
{ctx['ratios_text']}

{ctx['enhanced_text'] if ctx['enhanced_text'] else ''}

**模型分析数据（请翻译成大白话）**
{dupont}
{zscore}
{fscore}

**数据质量**
{ctx['quality_notes']}

**输出格式**：

## 一、体检结论（≥300字）
一段话总评你的企业整体财务健不健康。然后用红黄绿灯总览表展示5个维度（偿债/盈利/营运/现金流/税务）。

## 二、重点关注（≥1000字）
最多5个最需要注意的问题。每个问题用"四问"展开：
- 有没有问题？
- 严不严重？（和行业平均比）
- 什么原因？
- 不管会怎样？

包括模型体检结论（杜邦/Z-Score/F-Score各用1段大白话总结）。

## 三、行动清单（≥700字）
按优先级排列行动项。每项用"四答"展开：
- 做什么？
- 为什么？
- 怎么做？（给出具体步骤）
- 多久见效？

记住八条铁律：每个术语必须翻译、不用公式、红黄绿灯、四问四答、结论先行、像老朋友聊天。
"""

    # ═══════════════════════════════════════════════════════════
    # 模型解释固化文本
    # ═══════════════════════════════════════════════════════════

    def _build_model_explanations(self) -> str:
        """拼接三个模型的固化解释文本"""
        return (_MODEL_EXPLANATION_DUPONT + "\n\n" +
                _MODEL_EXPLANATION_ZSCORE + "\n\n" +
                _MODEL_EXPLANATION_FSCORE)

    # ═══════════════════════════════════════════════════════════
    # LLM调用
    # ═══════════════════════════════════════════════════════════

    def _call_llm(self, system_prompt: str, user_prompt: str, max_tokens: int = 4096) -> str:
        """调用LLM生成文本"""
        try:
            response = self._client.chat.completions.create(
                model=self._model,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt},
                ],
                temperature=0.7,
                max_tokens=max_tokens,
                timeout=120,
            )
            return response.choices[0].message.content
        except Exception:
            _logger.warning("LLM 调用失败，已脱敏处理")
            return "[模型调用失败，请稍后重试]"

    # ═══════════════════════════════════════════════════════════
    # 格式化工具
    # ═══════════════════════════════════════════════════════════

    @staticmethod
    def _fmt_value(v, unit=""):
        if isinstance(v, float) and isnan(v):
            return "N/A"
        if isinstance(v, float) and v == float("inf"):
            return "∞"
        if isinstance(v, dict):
            return str(v)
        if isinstance(v, float):
            if unit == "%":
                return f"{v*100:.2f}%"
            if abs(v) > 1e6:
                return f"{v:,.0f}"
            if abs(v) >= 1:
                return f"{v:.4f}"
            return f"{v:.6f}"
        return str(v)

    def _all_ratios_to_text(self, ratios: dict, benchmarks: dict) -> str:
        """所有指标格式化为文本（含行业基准）"""
        lines = []
        for category, items in ratios.items():
            lines.append(f"\n【{category}】")
            for name, info in items.items():
                v = info["value"]
                if isinstance(v, dict):
                    continue  # 跳过分变量dict
                if isinstance(v, float) and (isnan(v) or v == float("inf")):
                    continue
                unit = info.get("unit", "")
                bench = benchmarks.get(name, info.get("benchmark", ""))
                val_str = self._fmt_value(v, unit)
                bench_str = self._fmt_value(bench, unit) if isinstance(bench, (int, float)) else str(bench)
                lines.append(f"  {name}: {val_str} (行业基准: {bench_str})")
        return "\n".join(lines)

    def _enhanced_to_text(self, enhanced: dict, benchmarks: dict) -> str:
        """增强指标格式化为文本"""
        lines = []
        for dim, metrics in enhanced.items():
            lines.append(f"\n【{dim}·增强指标】")
            for name, info in metrics.items():
                v = info["value"]
                if isinstance(v, float) and (isnan(v) or v == float("inf")):
                    continue
                unit = info.get("unit", "")
                bench = benchmarks.get(name, "")
                val_str = self._fmt_value(v, unit)
                bench_str = self._fmt_value(bench, unit) if isinstance(bench, (int, float)) else str(bench)
                lines.append(f"  {name}: {val_str} (行业基准: {bench_str})")
        return "\n".join(lines)

    def _growth_to_text(self, growth: dict) -> str:
        """成长指标格式化"""
        lines = []
        for name, info in growth.items():
            v = info["value"]
            if isinstance(v, float) and isnan(v):
                lines.append(f"  {name}: 无法计算（数据不足）")
            else:
                lines.append(f"  {name}: {self._fmt_value(v, '%')}")
        return "\n".join(lines)

    def _benchmarks_to_text(self, benchmarks: dict) -> str:
        """行业基准格式化"""
        lines = []
        for name, val in benchmarks.items():
            if isinstance(val, float):
                if name in ("流动比率", "速动比率", "利息保障倍数", "产权比率", "现金比率", "净资产融资支撑度"):
                    lines.append(f"  {name}: {val:.2f}")
                elif name in ("总资产周转率", "固定资产周转率"):
                    lines.append(f"  {name}: {val:.2f}次")
                elif name in ("存货周转天数", "应收账款周转天数", "应付账款周转天数"):
                    lines.append(f"  {name}: {val:.0f}天")
                elif name == "Altman Z-score":
                    lines.append(f"  {name}: {val:.2f}")
                else:
                    lines.append(f"  {name}: {val*100:.2f}%")
            else:
                lines.append(f"  {name}: {val}")
        return "\n".join(lines)

    @staticmethod
    def _make_quality_notes(check_results) -> str:
        notes = []
        for r in check_results:
            if r.severity == "block":
                notes.append(f"[阻断] {r.check_name}: {r.detail}")
            elif r.severity == "warn":
                notes.append(f"[警告] {r.check_name}: {r.detail}")
            elif r.passed:
                notes.append(f"[通过] {r.check_name}: {r.detail}")
        return "\n".join(notes) if notes else "所有质量校验通过"


# ═══════════════════════════════════════════════════════════
# Markdown → docx 转换
# ═══════════════════════════════════════════════════════════

def markdown_to_docx(markdown_text: str, title: str = "财务分析报告") -> bytes:
    """将Markdown文本转换为docx文件（返回bytes）"""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei'
    font.size = Pt(11)

    # 标题
    h = doc.add_heading(title, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    lines = markdown_text.split('\n')
    in_table = False
    table_rows = []

    for line in lines:
        stripped = line.strip()

        # 表格处理
        if stripped.startswith('|') and stripped.endswith('|'):
            cells = [c.strip() for c in stripped.split('|')[1:-1]]
            if all(set(c) <= set('-: ') for c in cells):
                continue  # 分隔行
            table_rows.append(cells)
            in_table = True
            continue
        elif in_table:
            # 表格结束
            if table_rows:
                table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
                table.style = 'Table Grid'
                for i, row in enumerate(table_rows):
                    for j, cell in enumerate(row):
                        if j < len(table.columns):
                            table.cell(i, j).text = cell
                            if i == 0:
                                for p in table.cell(i, j).paragraphs:
                                    for r in p.runs:
                                        r.bold = True
            table_rows = []
            in_table = False

        # 标题
        if stripped.startswith('# ') and not stripped.startswith('## '):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith('## '):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith('### '):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith('#### '):
            doc.add_heading(stripped[5:], level=4)
        elif stripped.startswith('---'):
            doc.add_paragraph('─' * 50)
        elif stripped.startswith('━━'):
            doc.add_paragraph(stripped)
        elif stripped.startswith('• ') or stripped.startswith('- '):
            doc.add_paragraph(stripped[2:], style='List Bullet')
        elif stripped.startswith('① ') or stripped.startswith('② ') or stripped.startswith('③ '):
            doc.add_paragraph(stripped)
        elif stripped.startswith('**') and stripped.endswith('**'):
            p = doc.add_paragraph()
            run = p.add_run(stripped[2:-2])
            run.bold = True
        elif stripped == '':
            continue
        else:
            # 普通段落，处理加粗
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', stripped)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)

    # 处理剩余表格
    if table_rows:
        table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        table.style = 'Table Grid'
        for i, row in enumerate(table_rows):
            for j, cell in enumerate(row):
                if j < len(table.columns):
                    table.cell(i, j).text = cell

    # 保存到bytes
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def professional_report_to_docx(report: dict, title: str = "企业财务分析报告（专业版）") -> bytes:
    """将专业版报告dict转为docx"""
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei'
    font.size = Pt(11)

    h = doc.add_heading(title, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    module_names = [
        ("执行摘要", "一、执行摘要"),
        ("财务诊断", "二、财务诊断"),
        ("经典模型分析", "三、经典模型分析"),
        ("优化改进建议", "四、优化改进建议"),
        ("税务分析", "五、税务分析"),
        ("融资能力评估", "六、融资能力评估（Council多角色评估）"),
        ("免责声明", "七、免责与合规声明"),
    ]

    for key, title_text in module_names:
        content = report.get(key)
        if not content:
            continue
        doc.add_heading(title_text, level=1)
        # 逐段添加
        for para in content.split('\n'):
            stripped = para.strip()
            if not stripped:
                continue
            if stripped.startswith('### '):
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith('## '):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith('# '):
                doc.add_heading(stripped[2:], level=2)
            elif stripped.startswith('---') or stripped.startswith('━━'):
                doc.add_paragraph('─' * 50)
            elif stripped.startswith('• ') or stripped.startswith('- '):
                doc.add_paragraph(stripped[2:], style='List Bullet')
            else:
                p = doc.add_paragraph()
                parts = re.split(r'(\*\*.*?\*\*)', stripped)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    else:
                        p.add_run(part)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
